#!/usr/bin/env python3
"""
Convert docs/مستند-پروژه.md into a Word document with correct RTL layout,
Persian typography and thesis-style headings.

Usage:  python3 docs/build_docx.py
Output: docs/مستند-پروژه.docx
"""
import re
import pathlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

BODY_FONT = "B Nazanin"     # falls back gracefully if unavailable
LATIN_FONT = "Times New Roman"
MONO_FONT = "Consolas"

HERE = pathlib.Path(__file__).parent
SRC = HERE / "مستند-پروژه.md"
OUT = HERE / "مستند-پروژه.docx"


def set_rtl(paragraph):
    """Mark a paragraph as right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def style_run(run, size=12, bold=False, font=None, color=None):
    font_name = font or BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), LATIN_FONT if font != MONO_FONT else MONO_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT if font != MONO_FONT else MONO_FONT)
    rFonts.set(qn("w:cs"), font_name)
    # complex-script size/bold so Persian text honours them
    sz = OxmlElement("w:szCs")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement("w:bCs")
        rPr.append(b)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def add_par(doc, text="", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            font=None, space_after=6, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        style_run(p.add_run(text), size=size, bold=bold, font=font, color=color)
    return p


def add_heading(doc, text, level):
    sizes = {0: 20, 1: 16, 2: 14, 3: 12.5}
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(14 if level > 0 else 0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(text), size=sizes.get(level, 12), bold=True,
              color=RGBColor(0x1a, 0x1a, 0x1a))
    return p


def add_code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(line if line else " ")
        run.font.size = Pt(9.5)
        run.font.name = MONO_FONT
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), MONO_FONT)
        rFonts.set(qn("w:hAnsi"), MONO_FONT)
        rFonts.set(qn("w:cs"), MONO_FONT)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = 2  # right
    # right-to-left table
    tblPr = t._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

    for i, cell in enumerate(t.rows[0].cells):
        cell.text = ""
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(header[i]), size=10.5, bold=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EFEFEF")
        cell._tc.get_or_add_tcPr().append(shd)

    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            p = cells[i].paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            style_run(p.add_run(val), size=10.5)
    doc.add_paragraph()


def parse_table(lines, idx):
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            rows.append(cells)
        idx += 1
    return rows, idx


def clean(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.right_margin = Cm(3.0)   # gutter side for RTL binding
    sec.left_margin = Cm(2.5)
    # make the whole section RTL
    sectPr = sec._sectPr
    rtlGutter = OxmlElement("w:rtlGutter")
    sectPr.append(rtlGutter)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(12)

    lines = md.split("\n")
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_buf)
                doc.add_paragraph()
                code_buf, in_code = [], False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if 'page-break-after' in line:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            continue

        if stripped.startswith("#"):
            m = re.match(r"(#+)\s*(.+)", stripped)
            level = len(m.group(1)) - 1
            add_heading(doc, clean(m.group(2)), level)
            i += 1
            continue

        if stripped in ("---", "***", "<br>") or stripped.startswith("<div") \
           or stripped.startswith("</div") or stripped.startswith("<"):
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            txt = clean(re.sub(r"^[-*]\s+", "", stripped))
            if txt:
                p = add_par(doc, "• " + txt, size=11.5,
                            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=3)
                p.paragraph_format.right_indent = Cm(0.7)
            i += 1
            continue

        if re.match(r"^\d+[\.\)]\s+", stripped) or re.match(r"^[۰-۹]+[\.\)]\s+", stripped):
            txt = clean(stripped)
            p = add_par(doc, txt, size=11.5, align=WD_ALIGN_PARAGRAPH.RIGHT,
                        space_after=3)
            p.paragraph_format.right_indent = Cm(0.7)
            i += 1
            continue

        if stripped:
            txt = clean(stripped)
            if txt:
                bold = stripped.startswith("**") and stripped.rstrip().endswith("**")
                add_par(doc, txt, size=12, bold=bold)
        i += 1

    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
